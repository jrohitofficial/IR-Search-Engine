"""Reusable python-docx building blocks: shaded code panels, figures with
auto-numbered captions, tables, TODO callouts, and a real Word TOC field."""
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

NAVY = RGBColor(0x14, 0x1B, 0x3C)
VIOLET = RGBColor(0x5B, 0x3D, 0xF0)
TEAL = RGBColor(0x0F, 0x7E, 0xA3)
MUTED = RGBColor(0x5A, 0x6B, 0x7D)
AMBER_BG = "FFF4DC"
AMBER_BORDER = "E8A93C"
CODE_BG = "1E2530"
CODE_FG = "D7DEE9"

ROOT = Path(__file__).resolve().parents[1]

_fig_counter = [0]
_tab_counter = [0]


def next_fig_no():
    _fig_counter[0] += 1
    return _fig_counter[0]


def next_tab_no():
    _tab_counter[0] += 1
    return _tab_counter[0]


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, hex_color, sz=6):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), hex_color)
        borders.append(el)
    tcPr.append(borders)


def add_toc_field(document):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldText = OxmlElement("w:t")
    fldText.text = "Right-click and choose “Update Field” (or press F9) in Microsoft Word to generate the Table of Contents."
    fldChar2.append(fldText)
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    return paragraph


def add_title_page(document, title, subtitle, module_code, student_name, student_id, module_leader, date_str):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run(module_code)
    run.font.size = Pt(15)
    run.font.color.rgb = TEAL
    run.font.bold = True

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(title)
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = NAVY

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(13)
    run.font.italic = True
    run.font.color.rgb = MUTED

    document.add_paragraph().paragraph_format.space_before = Pt(40)

    table = document.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows_data = [
        ("Submitted by", student_name),
        ("Student ID", student_id),
        ("Submitted to", f"{module_leader} (Module Leader)"),
        ("Date", date_str),
    ]
    for i, (label, value) in enumerate(rows_data):
        c0, c1 = table.rows[i].cells
        c0.text = ""
        r = c0.paragraphs[0].add_run(label)
        r.font.bold = True
        r.font.size = Pt(11)
        c1.text = ""
        r = c1.paragraphs[0].add_run(value)
        r.font.size = Pt(11)
    document.add_page_break()


def add_h1(document, text):
    h = document.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h


def add_h2(document, text):
    h = document.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = VIOLET
    return h


def add_h3(document, text):
    h = document.add_heading(text, level=3)
    for run in h.runs:
        run.font.color.rgb = TEAL
    return h


def add_para(document, text, *, bold=False, italic=False, size=10.5, color=None):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.font.bold = bold
    run.font.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_bullets(document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(10.5)


def add_todo(document, prompt):
    table = document.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, AMBER_BG)
    set_cell_border(cell, AMBER_BORDER, sz=10)
    cell.text = ""
    p = cell.paragraphs[0]
    r1 = p.add_run("AUTHOR ANALYSIS REQUIRED — ")
    r1.font.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = RGBColor(0x8A, 0x5A, 0x00)
    r2 = p.add_run(prompt)
    r2.font.italic = True
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor(0x5A, 0x45, 0x10)
    document.add_paragraph()


def add_code(document, code_text, label=""):
    lines = code_text.rstrip("\n").split("\n")
    table = document.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, CODE_BG)
    set_cell_border(cell, "0D1117", sz=4)
    cell.text = ""
    tcPr = cell._tc.get_or_add_tcPr()
    for m in ("top", "bottom", "left", "right"):
        mar = OxmlElement(f"w:{m}")
        mar.set(qn("w:w"), "150")
        mar.set(qn("w:type"), "dxa")
        tcPr.append(mar)
    first = True
    if label:
        p = cell.paragraphs[0]
        r = p.add_run(label)
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x8B, 0x94, 0x9E)
        r.font.italic = True
        first = False
    for line in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        run = p.add_run(line if line.strip() else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8.7)
        run.font.color.rgb = RGBColor(0xD7, 0xDE, 0xE9)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
    document.add_paragraph()


def add_figure(document, image_path, caption, width_in=6.0):
    fig_no = next_fig_no()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Figure {fig_no}. ")
    r.font.bold = True
    r.font.size = Pt(9.5)
    r2 = cap.add_run(caption)
    r2.font.size = Pt(9.5)
    r2.font.italic = True
    r2.font.color.rgb = MUTED
    document.add_paragraph()
    return fig_no


def add_table(document, headers, rows, caption=None, col_widths=None):
    tab_no = next_tab_no()
    if caption:
        cap = document.add_paragraph()
        r = cap.add_run(f"Table {tab_no}. ")
        r.font.bold = True
        r.font.size = Pt(9.5)
        r2 = cap.add_run(caption)
        r2.font.size = Pt(9.5)
        r2.font.italic = True
        r2.font.color.rgb = MUTED

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr_cells[i], "1A3D6D")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9.3)
    document.add_paragraph()
    return tab_no


def read_source(relpath: str) -> str:
    return (ROOT / relpath).read_text(encoding="utf-8")


def extract_function(source: str, func_name: str) -> str:
    pattern = rf"^((?:async )?def {re.escape(func_name)}\(.*?)(?=\n(?:async )?def |\nclass |\Z)"
    m = re.search(pattern, source, re.MULTILINE | re.DOTALL)
    return m.group(1).rstrip() if m else f"# function {func_name} not found"
